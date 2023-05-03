from flask import Flask, render_template, request, send_file
from openpyxl import load_workbook
from docx import Document
from io import BytesIO
from uuid import uuid4
from weasyprint import HTML

app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        excel_file = request.files['excel_file']
        word_template = request.files['word_template']
        sheet_name = request.form['sheet_name']
        keys = [value for key, value in request.form.items() if key.startswith("key_")]

        # Read Excel data
        workbook = load_workbook(filename=BytesIO(excel_file.read()), data_only=True)
        sheet = workbook[sheet_name]
        values = [cell for row in sheet.iter_rows(min_row=2, max_row=2, values_only=True) for cell in row]

        # Read Word template and replace placeholders
        document = Document(BytesIO(word_template.read()))
        for key, value in zip(keys, values):
            for paragraph in document.paragraphs:
                if f'{{{key}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{key}}}', str(value) if value else '')

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if f'{{{key}}}' in cell.text:
                            cell.text = cell.text.replace(f'{{{key}}}', str(value) if value else '')

        # Save the modified Word document to memory
        output_file = BytesIO()
        document.save(output_file)
        output_file.seek(0)

        action = request.form.get("action")
        generate_pdf = action == "Generate PDF"

        if generate_pdf:
            output_pdf = BytesIO()
            # HTML(string=html_string).write_pdf(output_pdf)
            output_pdf.seek(0)
            mimetype = 'application/pdf'
            output_filename = f"{excel_file.filename[:-5]}_output.pdf"
        else:
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            output_filename = f"{excel_file.filename[:-5]}_output.docx"

        response = send_file(output_pdf if generate_pdf else output_file, mimetype=mimetype, as_attachment=True, download_name=output_filename)
        response.headers['Content-Disposition'] = f'attachment; filename={output_filename}'
        return response


    return render_template('index.html')


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get("PORT", 8080)), debug=True)
